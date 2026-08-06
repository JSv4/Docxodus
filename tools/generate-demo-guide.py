#!/usr/bin/env python3
"""Generate the branded DOCX used by the GitHub Pages editor demo.

Requires python-docx (`python -m pip install python-docx`). The generated file
is committed so the static demo itself has no build-time dependency.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "docs" / "demo" / "docxodus-demo-guide.docx"

NAVY = "071A2D"
NAVY_2 = "102A46"
BLUE = "245BDB"
CYAN = "13BBD4"
VIOLET = "7655D9"
GREEN = "159B72"
INK = "142238"
MUTED = "52657D"
PALE_BLUE = "EAF3FF"
PALE_CYAN = "E8FAFC"
PALE_VIOLET = "F1EDFF"
PALE_GREEN = "EAF8F2"
LINE = "CED9E7"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    for old in properties.findall(qn("w:shd")):
        properties.remove(old)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8, **edges) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    selected = edges or {"top": {}, "start": {}, "bottom": {}, "end": {}}
    for edge, options in selected.items():
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), options.get("val", "single"))
        node.set(qn("w:sz"), str(options.get("size", size)))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), options.get("color", color))


def remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                start={"val": "nil"},
                bottom={"val": "nil"},
                end={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def keep_with_next(paragraph_or_style) -> None:
    if hasattr(paragraph_or_style, "_p"):
        properties = paragraph_or_style._p.get_or_add_pPr()
    else:
        properties = paragraph_or_style.element.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    properties.append(keep)


def set_paragraph_border(paragraph, side: str, color: str, size=18, space=10) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def set_paragraph_shading(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, display, end))


def add_hyperlink(paragraph, label: str, url: str, color=BLUE):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((fonts, run_color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_run(paragraph, text: str, *, bold=False, color=INK, size=None, font="Arial", italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_kicker(container, text: str, color=CYAN):
    paragraph = container.add_paragraph(style="Kicker")
    add_run(paragraph, text.upper(), bold=True, color=color, size=9)
    return paragraph


def add_heading(container, text: str, level=1, number: str | None = None):
    if number:
        add_kicker(container, number)
    return container.add_paragraph(text, style=f"Heading {level}")


def add_body(container, text: str, *, bold_prefix: str | None = None, color=MUTED):
    paragraph = container.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        add_run(paragraph, bold_prefix, bold=True, color=INK)
        add_run(paragraph, text[len(bold_prefix):], color=color)
    else:
        add_run(paragraph, text, color=color)
    return paragraph


def add_pill_row(document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = (
        ("1 / SELECT", "Highlight text on this page", PALE_CYAN, CYAN),
        ("2 / EDIT", "Use the live ribbon above", PALE_BLUE, BLUE),
        ("3 / DOWNLOAD", "Open the result in Word", PALE_VIOLET, VIOLET),
    )
    for index, (title, copy, fill, accent) in enumerate(labels):
        cell = table.cell(0, index)
        cell.width = Inches(2.15)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=150, start=160, bottom=150, end=160)
        set_cell_border(cell, color="FFFFFF", size=16)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(3)
        add_run(paragraph, title, bold=True, color=accent, size=8)
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, copy, bold=True, color=INK, size=9.5)


def add_callout(document, label: str, title: str, copy: str, *, fill=PALE_BLUE, accent=BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.9)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    set_cell_border(cell, color=accent, size=10, start={"color": accent, "size": 28})
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    add_run(paragraph, label.upper(), bold=True, color=accent, size=8)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    add_run(paragraph, title, bold=True, color=INK, size=13)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, copy, color=MUTED, size=9.5)


def add_practice_paragraph(document, label: str, text: str, *, fill=PALE_VIOLET, accent=VIOLET):
    label_paragraph = document.add_paragraph()
    label_paragraph.paragraph_format.left_indent = Inches(.16)
    label_paragraph.paragraph_format.right_indent = Inches(.16)
    label_paragraph.paragraph_format.space_before = Pt(5)
    label_paragraph.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(label_paragraph, fill)
    set_paragraph_border(label_paragraph, "left", accent, size=26, space=10)
    add_run(label_paragraph, label.upper(), bold=True, color=accent, size=8)

    practice_paragraph = document.add_paragraph()
    practice_paragraph.paragraph_format.left_indent = Inches(.16)
    practice_paragraph.paragraph_format.right_indent = Inches(.16)
    practice_paragraph.paragraph_format.space_before = Pt(0)
    practice_paragraph.paragraph_format.space_after = Pt(5)
    set_paragraph_shading(practice_paragraph, fill)
    set_paragraph_border(practice_paragraph, "left", accent, size=26, space=10)
    add_run(practice_paragraph, text, color=INK, size=13)
    return practice_paragraph


def add_step_table(document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(.62), Inches(2.38), Inches(3.77))
    headers = ("STEP", "DO THIS", "YOU SHOULD SEE")
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        cell.width = width
        set_cell_shading(cell, NAVY_2)
        set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        set_cell_border(cell, color=NAVY_2)
        add_run(cell.paragraphs[0], label, bold=True, color=WHITE, size=8)
    set_repeat_table_header(table.rows[0])
    steps = (
        ("01", "Select the purple sentence", "A blue text selection appears."),
        ("02", "Click B on the Home tab", "The sentence becomes bold."),
        ("03", "Click Undo, then Redo", "The formatting disappears and returns."),
        ("04", "Click Save", "A real Word file saves to your device."),
    )
    for row_index, values in enumerate(steps):
        cells = table.add_row().cells
        fill = "F7FAFD" if row_index % 2 == 0 else WHITE
        for cell, width, value in zip(cells, widths, values):
            cell.width = width
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
            set_cell_border(cell, color=LINE, size=6)
            add_run(
                cell.paragraphs[0],
                value,
                bold=cell is cells[0] or cell is cells[1],
                color=BLUE if cell is cells[0] else INK,
                size=8.7 if cell is cells[1] else 9,
            )


def add_feature_matrix(document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("DOCUMENT FEATURE", "IN THE BROWSER", "IN THE SAVED DOCX")
    widths = (Inches(2.4), Inches(2.1), Inches(2.25))
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        cell.width = width
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        set_cell_border(cell, color=NAVY)
        add_run(cell.paragraphs[0], label, bold=True, color=WHITE, size=7.8)
    set_repeat_table_header(table.rows[0])
    rows = (
        ("Styles & inline formatting", "Faithful render + live edit", "Native run and style properties"),
        ("Tables & numbering", "Structured, selectable content", "Real tables and list definitions"),
        ("Footnotes & hyperlinks", "Linked references", "Relationship-backed OOXML"),
        ("Headers, footers & sections", "Full document context", "Original section structure"),
        ("Tracked changes", "Visible redline semantics", "Word-native revisions"),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = "F5F8FC" if index % 2 == 0 else WHITE
        for cell, width, value in zip(cells, widths, values):
            cell.width = width
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=105, start=130, bottom=105, end=130)
            set_cell_border(cell, color=LINE, size=6)
            add_run(cell.paragraphs[0], value, bold=cell is cells[0], color=INK if cell is cells[0] else MUTED, size=8.8)


def add_bullet(container, title: str, copy: str, color=BLUE):
    paragraph = container.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(5)
    add_run(paragraph, title, bold=True, color=color, size=10)
    add_run(paragraph, f" — {copy}", color=MUTED, size=10)
    return paragraph


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color, before, after in (
        ("Title", 34, WHITE, 0, 8),
        ("Subtitle", 14, "BED7F5", 0, 0),
        ("Heading 1", 23, NAVY, 13, 8),
        ("Heading 2", 14, NAVY_2, 11, 6),
        ("Heading 3", 11, BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    kicker = styles.add_style("Kicker", 1)
    kicker.font.name = "Arial"
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(CYAN)
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(3)
    keep_with_next(kicker)

    code = styles.add_style("Demo Code", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb("D7EBFF")
    code.paragraph_format.left_indent = Inches(.18)
    code.paragraph_format.right_indent = Inches(.18)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(.6)
    section.bottom_margin = Inches(.58)
    section.left_margin = Inches(.72)
    section.right_margin = Inches(.72)
    section.header_distance = Inches(.27)
    section.footer_distance = Inches(.28)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(paragraph, "DOCXODUS", bold=True, color=BLUE, size=8)
    add_run(paragraph, "  /  BROWSER DOCUMENT ENGINE", color=MUTED, size=7)
    set_paragraph_border(paragraph, "bottom", LINE, size=5, space=5)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, "DOCXODUS  •  LIVE PRODUCT GUIDE   |   ", color=MUTED, size=7.5)
    add_field(paragraph, "PAGE")


def add_cover(document: Document) -> None:
    hero = document.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero.autofit = False
    cell = hero.cell(0, 0)
    cell.width = Inches(6.9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=420, start=360, bottom=410, end=360)
    set_cell_border(cell, color=NAVY)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(10)
    add_run(paragraph, "LIVE DOCX  •  CLICK INTO THE PAGE", bold=True, color=CYAN, size=9)
    paragraph = cell.add_paragraph(style="Title")
    add_run(paragraph, "Edit this document.\nIt’s real.", bold=True, color=WHITE, size=34)
    paragraph = cell.add_paragraph(style="Subtitle")
    add_run(paragraph, "Select text. Change it. Save the same Word file.", color="BED7F5", size=14, italic=True)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    add_pill_row(document)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    add_kicker(document, "Try it now  •  20-second proof", BLUE)
    title = document.add_paragraph("Make the sentence below bold.", style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    add_body(
        document,
        "Drag across the purple sentence to select it, then click B on the Home tab above.",
    )

    add_practice_paragraph(
        document,
        "Select only this sentence",
        "This is a real Word document. Make this sentence bold.",
        fill=PALE_VIOLET,
        accent=VIOLET,
    )

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    add_kicker(document, "What to do", VIOLET)
    add_step_table(document)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    add_callout(
        document,
        "The proof",
        "Open the download in Word or LibreOffice.",
        "Your sentence is still bold because the browser edited the DOCX itself—not a screenshot, canvas, or HTML copy.",
        fill=PALE_CYAN,
        accent=CYAN,
    )


def add_controls_page(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "Now try the rest.", 1, "01 / GUIDED TUTORIAL")
    add_body(document, "Work straight down this page. Each exercise takes a few seconds and changes real WordprocessingML in the open document.")

    add_heading(document, "1. Mix inline styles", 2)
    add_practice_paragraph(
        document,
        "Practice line",
        "A real Word document is editable in this browser.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_body(document, "Select a few words. Try italic, underline, strike, superscript, subscript, and a different point size.")

    add_heading(document, "2. Change a whole paragraph", 2)
    add_practice_paragraph(
        document,
        "Click anywhere in this paragraph",
        "Center this paragraph, set it to 18 pt, then increase its indent.",
        fill=PALE_CYAN,
        accent=CYAN,
    )
    add_body(document, "Use the alignment and font-size menus, then the indent arrows. Undo each step if you want to start over.")

    add_heading(document, "3. Add real structure", 2)
    add_callout(
        document,
        "Put the caret at the end of this sentence",
        "Insert a table, rule, or footnote.",
        "Use ▦ for a 2 × 2 table, ― for a horizontal rule, or ¹ for a footnote. Each one becomes native document structure.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(document, "4. Make a list", 2)
    add_body(document, "Click this paragraph, then press •≡ for bullets or 1≡ for numbering. Use the indent arrows to change its level.")

    add_callout(
        document,
        "Safe to experiment",
        "Undo and Redo are real history.",
        "Make a mess. Reverse it. Restore it. The editor keeps the rest of the document stable while the active block changes.",
        fill=PALE_VIOLET,
        accent=VIOLET,
    )


def add_architecture_page(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "Open the download. Inspect the proof.", 1, "02 / WHAT SURVIVES")
    add_body(document, "The edited file is still a structured DOCX. Text stays text, tables stay tables, and document relationships remain document relationships.")

    add_heading(document, "What Docxodus preserves", 2)
    add_feature_matrix(document)

    add_heading(document, "What happened in the browser", 2)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cards = (
        ("1", "LOAD", "Fetch or open DOCX bytes", PALE_BLUE, BLUE),
        ("2", "EDIT", "Target stable document anchors", PALE_CYAN, CYAN),
        ("3", "SAVE", "Rebuild a lossless .docx", PALE_GREEN, GREEN),
    )
    for cell, (number, title, copy, fill, accent) in zip(table.rows[0].cells, cards):
        cell.width = Inches(2.25)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=190, start=170, bottom=190, end=170)
        set_cell_border(cell, color=WHITE, size=16)
        paragraph = cell.paragraphs[0]
        add_run(paragraph, number, bold=True, color=accent, size=18)
        paragraph = cell.add_paragraph()
        add_run(paragraph, title, bold=True, color=INK, size=10)
        paragraph = cell.add_paragraph()
        add_run(paragraph, copy, color=MUTED, size=8.5)

    add_heading(document, "What never leaves this tab", 2)
    add_bullet(document, "Your source document", "the browser reads it directly into memory.", GREEN)
    add_bullet(document, "Your edits", "formatting and structural mutations happen in the local session.", GREEN)
    add_bullet(document, "Your exported file", "Save produces bytes locally and the browser downloads them.", GREEN)

    add_callout(
        document,
        "Privacy by architecture",
        "No document server is hiding behind the demo.",
        "Network traffic loads the application runtime and this public sample. Your edits are never posted back to Docxodus or GitHub Pages.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(document, "See the real page", 2)
    paragraph = add_body(document, "Open the Layout tab and tick Page view. The continuous editor becomes four Word-style page boxes without reopening the file or losing your edits.")
    add_run(paragraph, " That round trip is the product.", bold=True, color=BLUE)


def add_embed_page(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "Take the editor with you.", 1, "03 / EMBED DOCXODUS")
    add_body(document, "The demo offers two deployment modes. Choose a hosted iframe for the fastest proof, or import the versioned module when Docxodus should feel native to your application.")

    add_heading(document, "A. Drop-in player", 2)
    add_body(document, "Use the 480 × 480 hosted player in any website that accepts iframe HTML.")
    code_table = document.add_table(rows=1, cols=1)
    code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = code_table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=180, start=210, bottom=180, end=210)
    set_cell_border(cell, color=NAVY)
    lines = (
        '<iframe src="https://jsv4.github.io/Docxodus/demo/player.html"',
        '        width="480" height="480"',
        '        sandbox="allow-scripts allow-same-origin allow-downloads">',
        '</iframe>',
    )
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.style = "Demo Code"
        add_run(paragraph, line, color="D7EBFF", size=8.5, font="Courier New")

    add_heading(document, "B. Native module", 2)
    add_body(document, "Import the pinned ESM bundle, then point createEditor at a container and any CORS-readable DOCX URL.")
    module_table = document.add_table(rows=1, cols=1)
    module_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = module_table.cell(0, 0)
    set_cell_shading(cell, NAVY_2)
    set_cell_margins(cell, top=180, start=210, bottom=180, end=210)
    set_cell_border(cell, color=NAVY_2)
    lines = (
        'import { createEditor } from',
        '  "https://cdn.jsdelivr.net/npm/docxodus@9.1.1/dist/embed.bundle.js";',
        '',
        'await createEditor("#editor", "/documents/example.docx");',
    )
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.style = "Demo Code"
        add_run(paragraph, line or " ", color="D7EBFF", size=8.5, font="Courier New")

    add_callout(
        document,
        "Social sharing",
        "Share the landing page; embed the player on sites you control.",
        "X and LinkedIn open the live demo from a link card. They do not execute arbitrary iframe HTML pasted into a post.",
        fill=PALE_VIOLET,
        accent=VIOLET,
    )

    add_heading(document, "Keep exploring", 2)
    paragraph = document.add_paragraph()
    add_hyperlink(paragraph, "Live demo", "https://jsv4.github.io/Docxodus/demo/")
    add_run(paragraph, "   •   ", color=LINE)
    add_hyperlink(paragraph, "GitHub repository", "https://github.com/JSv4/Docxodus")
    add_run(paragraph, "   •   ", color=LINE)
    add_hyperlink(paragraph, "npm package", "https://www.npmjs.com/package/docxodus")

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(20)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(closing, "DOCX IN. DOCX OUT.", bold=True, color=CYAN, size=10)
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(closing, "The browser is the document workspace now.", bold=True, color=NAVY, size=18)


def build(output: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)
    add_controls_page(document)
    add_architecture_page(document)
    add_embed_page(document)

    properties = document.core_properties
    properties.title = "Edit This Document — Docxodus Live Tutorial"
    properties.subject = "Hands-on tutorial for the Docxodus browser DOCX editor"
    properties.author = "Docxodus"
    properties.keywords = "Docxodus, DOCX, OOXML, WebAssembly, browser editor"
    properties.comments = "Generated by tools/generate-demo-guide.py"

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Generated {output.relative_to(ROOT)} ({output.stat().st_size:,} bytes)")


if __name__ == "__main__":
    target = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_OUTPUT
    build(target)
